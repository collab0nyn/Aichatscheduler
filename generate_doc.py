from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(41, 128, 185) # Blue color

def create_document():
    doc = Document()

    # Title
    title = doc.add_heading('TaskSense: NLP-Based Smart Action Item Extractor', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Implementation Plan and Technical Approach').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 1. Problem Statement
    add_heading(doc, '1. Problem Statement', 1)
    doc.add_paragraph(
        "Students and teams often discuss tasks in WhatsApp/Discord/meeting notes, but action items get lost in long conversations. "
        "TaskSense is a full-stack NLP-based application that takes meeting text or chat messages as input and automatically identifies: "
        "Who has to do what, by when, and priority level."
    )

    # 2. Features
    add_heading(doc, '2. Key Features', 1)
    features = [
        "Login for teams",
        "Upload .txt meeting notes",
        "Paste WhatsApp/Discord chat",
        "Auto-generate task board with Tasks, Deadlines, Responsible People, and Priorities",
        "Mark tasks as completed",
        "Send reminders",
        "Export action plan to PDF/CSV"
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    # 3. Technical Approach (Full Stack)
    add_heading(doc, '3. Technical Approach', 1)
    
    add_heading(doc, 'Frontend (React.js + Vanilla CSS)', 2)
    doc.add_paragraph(
        "The frontend is built using React (via Vite) for a fast, responsive user interface. "
        "It features a premium, dynamic design with dark mode, glassmorphism, and smooth animations using pure CSS. "
        "Lucide React is used for crisp iconography. Key components include a Dashboard, an Input Section for text/file upload, "
        "and a Kanban-style Task Board."
    )

    add_heading(doc, 'Backend (Flask / Node.js)', 2)
    doc.add_paragraph(
        "The backend will serve as a REST API. A Python Flask backend is recommended because it integrates seamlessly with Python NLP libraries."
    )

    add_heading(doc, 'Database (Supabase / MongoDB)', 2)
    doc.add_paragraph(
        "Data persistence for users, teams, and generated task boards will be handled by a modern database like Supabase (PostgreSQL) or MongoDB."
    )

    add_heading(doc, 'NLP Engine (Python Rule-Based)', 2)
    doc.add_paragraph(
        "The core extraction engine will use simple, lightweight NLP techniques without the need for model training:"
    )
    nlp_techniques = [
        "Regex: For extracting dates and deadlines.",
        "Keyword Matching: For priority levels (e.g., 'urgent', 'ASAP', 'important').",
        "Rule-based Extraction: Using action verbs ('do', 'submit', 'prepare', 'complete').",
        "Named Entity Extraction: Using spaCy (optional) to identify people's names.",
        "Text Summarization: Using a simple sentence ranking algorithm or external API to generate meeting summaries."
    ]
    for tech in nlp_techniques:
        doc.add_paragraph(tech, style='List Bullet')

    # 4. Step-by-Step Solution Workflow
    add_heading(doc, '4. Step-by-Step Implementation Workflow', 1)
    
    steps = [
        ("Step 1: Frontend Initialization", "Set up the React Vite project and establish the global design system (tokens, colors, animations)."),
        ("Step 2: UI Component Development", "Build the layout, InputSection for pasting chat/uploading notes, and the TaskBoard for displaying extracted data."),
        ("Step 3: Backend API Setup", "Initialize a Flask application with endpoints for /extract and /summarize."),
        ("Step 4: NLP Logic Implementation", "Write Python scripts using regex and keyword matching to process the text payload and return structured JSON."),
        ("Step 5: Database Integration", "Connect to Supabase/MongoDB to save users, teams, and action plans."),
        ("Step 6: Final Integration & Export", "Connect the React frontend to the Flask backend, and implement PDF/CSV export functionality.")
    ]

    for title, desc in steps:
        p = doc.add_paragraph()
        p.add_run(title + ": ").bold = True
        p.add_run(desc)

    # Save
    doc.save('TaskSense_Technical_Approach.docx')
    print("Word document generated successfully: TaskSense_Technical_Approach.docx")

if __name__ == '__main__':
    create_document()
